from __future__ import annotations

import json
import re
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


RESULT_DIR = Path("results/2026-05-31T04-29-10-054Z-help-inventory")
INPUT_JSON = RESULT_DIR / "help-inventory-pilot.json"
OUTPUT_DOCX = RESULT_DIR / "Common Ground Help Inventory Pilot.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def clean_text(value: str | None) -> str:
    text = re.sub(r"\s+", " ", value or "").strip()
    if not text:
        return ""
    if re.fullmatch(r"/[A-Za-z0-9_/?=&.-]+", text):
        return "Dashboard icon/link" if text == "/dashboard" else text
    if re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", text):
        return "Support email link" if text.lower().startswith("support@") else "Email field"
    if text == "John Stevens":
        return "Name field"
    return text


def element_name(item: dict) -> str:
    return clean_text(
        item.get("label")
        or item.get("text")
        or item.get("placeholder")
        or item.get("ariaLabel")
        or item.get("href")
        or item.get("selector")
        or item.get("tag")
    )


def section_name(item: dict) -> str:
    section = clean_text(item.get("headingContext"))
    if re.fullmatch(r"CG-\d+ .+", section):
        return "Case card"
    return section or "Page navigation / main content"


def concise_guidance(item: dict) -> str:
    name = element_name(item)
    text = item.get("draftHelpDirection") or ""
    if "Explain what" in text:
        lower = name.lower()
        if lower == "getting started":
            return "Explain that this opens the first discussion step and helps capture the user's initial perspective."
        if lower == "case details":
            return "Explain that this opens the case overview, current status, and available next actions."
        if lower == "create new case":
            return "Explain when to start a new case and what information the user should have ready."
        if lower == "reports":
            return "Explain that reports summarize completed analysis and next-step guidance."
        if lower.startswith("step:"):
            return "Explain the current workflow step and what the user can review from here."
        if lower == "view":
            return "Explain what record or detail panel will open before the user selects it."
        if lower == "agreement readiness by term":
            return "Explain how readiness helps users see which discussion terms may be easier to resolve."
        if lower == "term-criterion alignment matrix":
            return "Explain how the matrix compares terms against decision criteria."
        if lower == "fact matrix snapshot":
            return "Explain that this section summarizes key facts used in the alignment analysis."
        if lower == "recommended discussion outline":
            return "Explain that this outline helps structure the follow-up conversation."
        if lower == "preparation for your conversation":
            return "Explain how to use the preparation guidance before meeting with the other party."
        if lower == "notifications 4" or "notifications" in lower:
            return "Explain that notifications call attention to updates or required user action."
        if lower == "go back" or lower == "back to case":
            return "Explain where the user will return and whether any unsaved changes are affected."
        if lower in {"account", "profile", "security", "dashboard", "contact us", "log-out"}:
            return f"Explain that this opens the {name} area of Common Ground."
        return f"Explain what {name} does and when the user should use it."
    return clean_text(text)


def dedupe_items(items: list[dict]) -> list[dict]:
    grouped: dict[tuple[str, str, str], dict] = {}
    counts: Counter[tuple[str, str, str]] = Counter()
    examples: defaultdict[tuple[str, str, str], set[str]] = defaultdict(set)
    priority_rank = {"High": 0, "Medium": 1, "Low": 2}

    for item in items:
        name = element_name(item)
        if not name:
            continue
        section = section_name(item)
        key = (name, item.get("kind", ""), section)
        counts[key] += 1
        if item.get("headingContext"):
            examples[key].add(clean_text(item.get("headingContext")))
        current = grouped.get(key)
        if current is None or priority_rank.get(item.get("priority", "Low"), 9) < priority_rank.get(current.get("priority", "Low"), 9):
            grouped[key] = item

    result = []
    for key, item in grouped.items():
        item = dict(item)
        item["_name"] = key[0]
        item["_section"] = key[2]
        item["_count"] = counts[key]
        sample = sorted(examples[key])[:3]
        item["_examples"] = ", ".join(sample)
        result.append(item)

    return sorted(result, key=lambda item: ({"High": 0, "Medium": 1, "Low": 2}.get(item.get("priority", "Low"), 3), item["_section"], item["_name"]))


def add_table(document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table)
    set_cell_margins(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    document.add_paragraph()
    return table


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def main() -> None:
    audits = json.loads(INPUT_JSON.read_text(encoding="utf-8"))
    document = Document()
    configure_styles(document)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Common Ground Help Inventory Pilot")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    run.bold = True

    subtitle = document.add_paragraph()
    subtitle.add_run("Requestor account, production read-only audit").italic = True
    document.add_paragraph(f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")

    document.add_heading("Summary", level=1)
    document.add_paragraph(
        "This document converts the read-only pilot capture into an implementation-oriented help inventory. "
        "The Dashboard has been deduplicated because the same controls repeat across many case cards."
    )
    document.add_paragraph(
        "Recommended next step: validate these help targets with product owners, then repeat the same capture for "
        "Participant-side cases and staged workflow pages that are not visible from this Requestor case."
    )

    coverage_rows = []
    for audit in audits:
        deduped = dedupe_items(audit["items"])
        coverage_rows.append([
            audit["pageName"],
            str(len(audit["items"])),
            str(len(deduped)),
            ", ".join(audit.get("headings", [])[:3]) or "No heading captured",
        ])
    add_table(
        document,
        ["Page", "Raw items", "Deduped targets", "Primary headings observed"],
        coverage_rows,
        [1.8, 0.8, 1.0, 3.8],
    )

    document.add_heading("Help Inventory", level=1)
    for audit in audits:
        document.add_heading(audit["pageName"], level=2)
        url_para = document.add_paragraph("URL: ")
        add_hyperlink(url_para, audit["url"], audit["url"])

        rows = []
        for item in dedupe_items(audit["items"]):
            count_note = f"Repeated {item['_count']} times" if item["_count"] > 1 else "Single instance"
            rows.append([
                item["_name"],
                item.get("kind", ""),
                item["_section"],
                item.get("recommendedHelpType", ""),
                item.get("priority", ""),
                count_note,
                concise_guidance(item),
            ])
        add_table(
            document,
            ["Element", "Type", "Section", "Help type", "Priority", "Frequency", "Help direction"],
            rows,
            [1.15, 0.55, 1.1, 0.95, 0.6, 0.8, 2.35],
        )

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Access Gaps For Full Coverage", level=1)
    gaps = [
        "Participant login and Participant dashboard state.",
        "Getting Started Discussion page with response input and inline coaching placement.",
        "Fact Label page.",
        "Emotion Moderation page.",
        "Cases paused at stages that are waiting on the Requestor versus waiting on the Participant.",
        "Any completed, closed, empty, error, mobile, or offline PWA states that differ from this pilot capture.",
    ]
    for gap in gaps:
        document.add_paragraph(gap, style="List Bullet")

    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
